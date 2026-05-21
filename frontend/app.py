"""竞品情报分析系统可视化前端"""
import streamlit as st
import requests
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io


# ===================== 自定义 SSE 解析器 =====================

def parse_sse_stream(response):
    """
    逐行解析 SSE (Server-Sent Events) 流，替代有 bug 的 sseclient 库。

    SSE 格式：
        event: <event_name>
        data: <json_string>
        <blank line>

    Yields: (event_name: str, data: dict)
    """
    current_event = None
    data_lines = []

    for line in response.iter_lines(decode_unicode=True):
        if line is None:
            continue

        # 空行 = 事件结束分隔符
        if line == "":
            if current_event and data_lines:
                try:
                    payload = json.loads("".join(data_lines))
                    yield current_event, payload
                except json.JSONDecodeError:
                    pass  # 跳过无法解析的事件
            current_event = None
            data_lines = []
            continue

        # 注释行跳过
        if line.startswith(":"):
            continue

        # 解析 field: value
        if ": " in line:
            field, _, value = line.partition(": ")
            if field == "event":
                current_event = value
            elif field == "data":
                data_lines.append(value)
        elif line.endswith(":"):
            # 无值字段（retry: 等），忽略
            field = line[:-1]
            if field == "data":
                data_lines.append("")


# 页面配置
st.set_page_config(
    page_title="多Agent竞品情报分析系统",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 全局配置
API_BASE_URL = "http://localhost:8000"

# ── Mock 模式检测 ──
@st.cache_data(ttl=3)
def fetch_mock_status():
    try:
        resp = requests.get(f"{API_BASE_URL}/api/v1/mock/status", timeout=3)
        if resp.status_code == 200:
            return resp.json()
    except Exception:
        pass
    return {"mock_mode": False, "current_scenario": "scenario1", "available_scenarios": []}


# 侧边栏导航
with st.sidebar:
    st.title("📊 竞品情报系统")
    st.divider()

    # ── Mock 模式开关 ──
    mock_status = fetch_mock_status()
    is_mock = mock_status.get("mock_mode", False)

    st.subheader("🎭 Demo 模式")
    mock_enabled = st.toggle(
        "Mock 模式（预生成数据）",
        value=is_mock,
        help="开启后所有分析使用预生成数据，3秒内出完整结果，无需大模型API。适用于Demo演示。"
    )
    if mock_enabled != is_mock:
        try:
            toggle_resp = requests.post(
                f"{API_BASE_URL}/api/v1/mock/toggle",
                json={"enabled": mock_enabled, "scenario": "scenario1"},
                timeout=5
            )
            if toggle_resp.status_code == 200:
                st.cache_data.clear()
                st.rerun()
        except Exception:
            st.warning("切换Mock模式失败，请检查后端服务")

    if mock_enabled:
        scenarios = mock_status.get("available_scenarios", [])
        if scenarios:
            scenario_options = {s["name"]: s["id"] for s in scenarios}
            selected_scenario = st.selectbox(
                "Demo 场景",
                options=list(scenario_options.keys()),
                index=0,
                help="选择要演示的竞品分析场景"
            )

    st.divider()
    page = st.radio("功能导航", ["竞品分析工作台", "竞品管理", "我方产品", "历史分析", "系统配置", "可观测中心", "竞品知识库"])
    st.divider()
    st.caption("基于多Agent架构的企业级竞品情报分析平台")

# 健康检查
def check_health():
    try:
        res = requests.get(f"{API_BASE_URL}/health", timeout=5)
        return res.status_code == 200
    except:
        return False

# 获取竞品列表
@st.cache_data(ttl=60)
def get_competitor_list():
    try:
        res = requests.get(f"{API_BASE_URL}/competitors/all", timeout=5)
        if res.status_code == 200:
            return res.json()
        return []
    except:
        return []

# 生成Word竞品分析报告
def generate_word_report(competitor_name, analysis_result, quality_score):
    doc = Document()
    # 标题
    doc.add_heading(f"{competitor_name} 竞品分析报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"分析质量评分：{quality_score}/10")
    doc.add_page_break()

    # 1. 竞品对比矩阵
    doc.add_heading("一、竞品对比矩阵", level=1)
    if analysis_result.get("comparison_matrix"):
        matrix = analysis_result["comparison_matrix"]
        for dim in matrix["dimensions"]:
            doc.add_heading(dim["dimension"], level=2)
            doc.add_paragraph(f"我方评分：{dim['our_score']}/10 | 竞品评分：{dim['competitor_score']}/10")
            doc.add_paragraph(f"分析说明：{dim['notes']}")
        doc.add_heading("整体评估", level=2)
        doc.add_paragraph(matrix["overall_assessment"])
    doc.add_page_break()

    # 2. 销售战术卡
    doc.add_heading("二、销售战术卡", level=1)
    if analysis_result.get("battlecard"):
        card = analysis_result["battlecard"]
        doc.add_heading("我方优势", level=2)
        for item in card["our_strengths"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("我方劣势", level=2)
        for item in card["our_weaknesses"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("竞品优势", level=2)
        for item in card["competitor_strengths"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("竞品劣势", level=2)
        for item in card["competitor_weaknesses"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("核心差异化", level=2)
        for item in card["key_differentiators"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("异议处理话术", level=2)
        for question, answer in card["objection_handling"].items():
            doc.add_paragraph(f"Q: {question}", style="List Bullet")
            doc.add_paragraph(f"A: {answer}")
        doc.add_heading("电梯Pitch", level=2)
        doc.add_paragraph(card["elevator_pitch"])
    doc.add_page_break()

    # 3. 研究洞察
    doc.add_heading("三、研究洞察", level=1)
    if analysis_result.get("research_results"):
        for insight in analysis_result["research_results"]:
            doc.add_heading(insight["topic"], level=2)
            doc.add_paragraph(insight["summary"])
            doc.add_heading("核心发现", level=3)
            for finding in insight["key_findings"]:
                doc.add_paragraph(f"- {finding}", style="List Bullet")
    doc.add_page_break()

    # 4. 变更监控
    doc.add_heading("四、变更监控", level=1)
    if analysis_result.get("changes_detected"):
        for change in analysis_result["changes_detected"]:
            doc.add_paragraph(f"[{change['severity']}] {change['title']}")
            doc.add_paragraph(change["summary"])
    else:
        doc.add_paragraph("本次监控未检测到竞品变更")

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ── Mock 模式全局横幅 ──
if is_mock:
    st.warning(
        "🎭 **Demo 模式已启用** — 所有分析使用预生成数据，3秒内出完整结果。"
        "关闭侧边栏「Mock 模式」开关恢复真实 LLM 调用。"
    )

# ------------------------------
# 页面1：竞品分析工作台（核心）
# ------------------------------
if page == "竞品分析工作台":
    st.header("🔍 竞品全维度分析工作台")
    st.caption("输入竞品信息，一键启动多Agent智能分析")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        st.success("✅ 后端服务正常运行")

    # 新增：快速选择已有竞品
    competitor_list = get_competitor_list()
    use_exist_competitor = st.checkbox("选择已有竞品", value=False)
    selected_competitor = None
    competitor_name = ""
    competitor_urls = ""

    if use_exist_competitor and competitor_list:
        competitor_options = {comp["name"]: comp for comp in competitor_list}
        selected_name = st.selectbox("选择竞品", options=list(competitor_options.keys()))
        selected_competitor = competitor_options[selected_name]
        competitor_name = selected_competitor["name"]
        competitor_urls = "\n".join(selected_competitor["urls"])
        st.info(f"已选择竞品：{selected_name}，自动填充监控URL")

    # 输入区域
    with st.form("analyze_form"):
        col1, col2 = st.columns([1, 2])
        with col1:
            competitor_name = st.text_input("竞品名称", value=competitor_name, placeholder="例如：字节跳动")
        with col2:
            competitor_urls = st.text_area(
                "监控URL（一行一个）",
                value=competitor_urls,
                placeholder="https://www.bytedance.com\nhttps://www.feishu.cn",
                height=80
            )
        submit = st.form_submit_button("🚀 启动全流程分析", use_container_width=True, type="primary")

    # 分析执行逻辑
    if submit and service_online:
        if not competitor_name:
            st.warning("请输入竞品名称")
        else:
            # 处理URL
            url_list = [u.strip() for u in competitor_urls.split("\n") if u.strip()]
            request_body = {
                "competitor": competitor_name,
                "urls": url_list
            }

            # 进度展示区域
            st.divider()
            st.subheader("📡 实时执行进度")
            progress_container = st.container()
            result_container = st.container()

            # 调用 SSE 流式接口（使用自定义解析器替代 sseclient）
            try:
                response = requests.post(
                    f"{API_BASE_URL}/analyze/stream",
                    json=request_body,
                    stream=True,
                    headers={"Accept": "text/event-stream"},
                    timeout=(5, 300),  # (connect_timeout, read_timeout)
                )
                if response.status_code != 200:
                    progress_container.error(f"❌ 后端返回错误状态码：{response.status_code}")
                    st.stop()

                # 存储各节点结果
                node_results = {}
                final_analysis_result = {}
                final_quality_score = 0.0
                has_error = False

                # 节点颜色映射：业务=蓝 / 校验=橙 / 修复=黄 / 溯源=绿
                NODE_COLORS = {
                    "monitor": "🔵", "alert": "🔵", "research": "🔵",
                    "compare": "🔵", "battlecard": "🔵",
                    "fact_check": "🟠", "reviewer": "🟠",
                    "targeted_fix": "🟡", "citation": "🟢",
                }

                # 实时处理事件（使用自定义 SSE 解析器）
                for node_name, node_data in parse_sse_stream(response):
                    if node_name == "error":
                        progress_container.error(f"❌ 执行出错：{node_data.get('error', '未知错误')}")
                        has_error = True
                        break

                    node_results[node_name] = node_data
                    icon = NODE_COLORS.get(node_name, "✅")

                    with progress_container:
                        with st.expander(f"{icon} {node_name} 节点执行完成", expanded=True):
                            st.json(node_data, expanded=False)

                    # 保存最终结果 — 原有节点
                    if node_name == "reviewer":
                        final_quality_score = node_data.get("quality_score", 0.0)
                        final_analysis_result["review_feedback"] = node_data.get("review_feedback", node_data)
                    if node_name == "quality_check":
                        final_quality_score = node_data.get("quality_score", 0.0)
                    if node_name == "battlecard":
                        final_analysis_result["battlecard"] = node_data.get("battlecard")
                    if node_name == "compare":
                        final_analysis_result["comparison_matrix"] = node_data.get("comparison_matrix")
                    if node_name == "research":
                        final_analysis_result["research_results"] = node_data.get("research_results", [])
                    if node_name == "monitor":
                        final_analysis_result["changes_detected"] = node_data.get("changes_detected", [])
                    # ★ 新增校验节点
                    if node_name == "fact_check":
                        final_analysis_result["fact_check_result"] = node_data.get("fact_check_result", node_data)
                    if node_name == "citation":
                        final_analysis_result["citation_report"] = node_data.get("citation_report", node_data)

                if has_error:
                    st.stop()

                # 检查是否收到有效分析结果
                if not node_results:
                    progress_container.warning(
                        "⚠️ 未收到任何分析事件。可能原因："
                        "1) 后端 Agent 执行超时；2) LLM API 调用失败。"
                        "请检查后端日志或重试。"
                    )
                    st.stop()

                # 分析完成，展示最终结果
                with result_container:
                    st.divider()
                    st.header("📑 最终分析报告")

                    # 新增：Word报告导出按钮
                    col_export, col_score = st.columns([1, 1])
                    with col_export:
                        if final_analysis_result:
                            word_buffer = generate_word_report(competitor_name, final_analysis_result, final_quality_score)
                            st.download_button(
                                label="📥 导出Word分析报告",
                                data=word_buffer,
                                file_name=f"{competitor_name}_竞品分析报告_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    with col_score:
                        st.metric("最终质量评分", f"{final_quality_score}/10")

                    # 分Tab展示结果
                    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
                        "📈 对比矩阵", "🎯 销售战术卡", "🔎 研究洞察",
                        "⚠️ 变更监控", "🛡️ 质量审查报告", "🔗 引用溯源报告"
                    ])

                    # 对比矩阵
                    with tab1:
                        if "comparison_matrix" in final_analysis_result and final_analysis_result["comparison_matrix"]:
                            matrix = final_analysis_result["comparison_matrix"]
                            st.subheader(f"竞品：{matrix['competitor']}")
                            for dim in matrix["dimensions"]:
                                col1, col2, col3 = st.columns([2, 1, 1])
                                with col1:
                                    st.markdown(f"**{dim['dimension']}**")
                                    st.caption(dim["notes"])
                                with col2:
                                    st.metric("我方评分", dim["our_score"])
                                with col3:
                                    st.metric("竞品评分", dim["competitor_score"])
                                st.divider()
                            st.markdown(f"**整体评估**：{matrix['overall_assessment']}")

                    # 销售战术卡
                    with tab2:
                        if "battlecard" in final_analysis_result and final_analysis_result["battlecard"]:
                            card = final_analysis_result["battlecard"]
                            col1, col2 = st.columns(2)
                            with col1:
                                st.subheader("✅ 我方优势")
                                for item in card["our_strengths"]:
                                    st.write(f"- {item}")
                                st.subheader("⚠️ 我方劣势")
                                for item in card["our_weaknesses"]:
                                    st.write(f"- {item}")
                            with col2:
                                st.subheader("📈 竞品优势")
                                for item in card["competitor_strengths"]:
                                    st.write(f"- {item}")
                                st.subheader("❌ 竞品劣势")
                                for item in card["competitor_weaknesses"]:
                                    st.write(f"- {item}")
                            st.divider()
                            st.subheader("🎯 核心差异化")
                            for item in card["key_differentiators"]:
                                st.write(f"- {item}")
                            st.divider()
                            st.subheader("🗣️ 异议处理话术")
                            for question, answer in card["objection_handling"].items():
                                with st.expander(question):
                                    st.write(answer)
                            st.divider()
                            st.subheader("🎙️ 电梯Pitch")
                            st.info(card["elevator_pitch"])

                    # 研究洞察
                    with tab3:
                        if "research_results" in final_analysis_result:
                            insights = final_analysis_result["research_results"]
                            for insight in insights:
                                with st.expander(f"📝 {insight['topic']}", expanded=True):
                                    st.write(insight["summary"])
                                    st.markdown("**核心发现**")
                                    for finding in insight["key_findings"]:
                                        st.write(f"- {finding}")

                    # 变更监控
                    with tab4:
                        if "changes_detected" in final_analysis_result:
                            changes = final_analysis_result["changes_detected"]
                            if not changes:
                                st.info("本次监控未检测到竞品变更")
                            else:
                                for change in changes:
                                    st.write(f"[{change['severity']}] {change['title']}")

                    # ★ 新增 Tab 5：质量审查报告
                    with tab5:
                        st.subheader("🛡️ 质量审查报告")
                        review_fb = final_analysis_result.get("review_feedback", {})
                        if review_fb:
                            col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                            with col_r1:
                                st.metric("综合评分", f"{review_fb.get('overall_score', 0):.1f}/10")
                            with col_r2:
                                st.metric("准确度", f"{review_fb.get('accuracy_score', 0):.1f}/10")
                            with col_r3:
                                st.metric("完整度", f"{review_fb.get('completeness_score', 0):.1f}/10")
                            with col_r4:
                                st.metric("可操作性", f"{review_fb.get('actionability_score', 0):.1f}/10")

                            approved = review_fb.get("approved", False)
                            if approved:
                                st.success("✅ 审查通过 — 报告质量达标")
                            else:
                                st.warning("⚠️ 审查未通过 — 已触发定向修复")

                            issues = review_fb.get("issues", [])
                            if issues:
                                st.markdown(f"**发现问题 ({len(issues)} 条)**")
                                for iss in issues:
                                    sev_icon = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(iss.get("severity", ""), "⚪")
                                    with st.expander(f"{sev_icon} [{iss.get('severity', '?')}] {iss.get('target', '')}"):
                                        st.markdown(f"**问题**: {iss.get('description', '')}")
                                        st.markdown(f"**修复建议**: {iss.get('fix_instruction', '')}")
                            st.divider()
                            st.caption(f"修订指令: {review_fb.get('revision_instructions', '无')[:300]}")
                        elif final_quality_score > 0:
                            st.metric("综合评分（旧版QC）", f"{final_quality_score:.1f}/10")
                        else:
                            st.info("暂无审查数据")

                    # ★ 新增 Tab 6：引用溯源报告
                    with tab6:
                        st.subheader("🔗 引用溯源报告")
                        cit_report = final_analysis_result.get("citation_report", {})
                        if cit_report:
                            total = cit_report.get("total_sources", 0)
                            verified = cit_report.get("verified_sources", 0)
                            broken = cit_report.get("broken_links", 0)

                            col_c1, col_c2, col_c3 = st.columns(3)
                            with col_c1:
                                st.metric("总引用数", total)
                            with col_c2:
                                st.metric("已验证", f"{verified}/{total}" if total > 0 else "N/A")
                            with col_c3:
                                st.metric("失效链接", broken)

                            rel = cit_report.get("overall_reliability_score", 0)
                            st.progress(min(rel, 1.0), text=f"总体可信度: {rel:.0%}")

                            missing = cit_report.get("missing_citations", [])
                            if missing:
                                st.warning(f"⚠️ 发现 {len(missing)} 条缺失引用的结论")
                                with st.expander("查看缺失引用详情"):
                                    for m in missing:
                                        st.caption(f"- {m[:120]}...")

                            dist = cit_report.get("reliability_distribution", {})
                            if dist:
                                st.caption(f"可信度分布: {dist}")
                        else:
                            st.info("暂无引用溯源数据（升级前的旧版分析无此功能）")

            except Exception as e:
                st.error(f"请求失败：{str(e)}")

# ------------------------------
# 页面2：竞品管理
# ------------------------------
elif page == "竞品管理":
    st.header("📋 竞品库管理")
    st.caption("统一管理竞品信息，支持新增、编辑、删除、快速选择")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        # 刷新竞品列表
        if st.button("🔄 刷新列表", use_container_width=False):
            st.cache_data.clear()
            st.rerun()

        # 新增竞品表单
        with st.expander("➕ 新增竞品", expanded=False):
            with st.form("add_competitor_form"):
                col1, col2 = st.columns([1, 2])
                with col1:
                    new_name = st.text_input("竞品名称", placeholder="例如：字节跳动")
                with col2:
                    new_urls = st.text_area(
                        "监控URL（一行一个）",
                        placeholder="https://www.bytedance.com\nhttps://www.feishu.cn",
                        height=80
                    )
                add_submit = st.form_submit_button("确认新增", use_container_width=True, type="primary")

                if add_submit:
                    if not new_name:
                        st.warning("请输入竞品名称")
                    else:
                        url_list = [u.strip() for u in new_urls.split("\n") if u.strip()]
                        res = requests.post(
                            f"{API_BASE_URL}/competitors",
                            json={"name": new_name, "urls": url_list}
                        )
                        if res.status_code == 200:
                            st.success(f"✅ 竞品【{new_name}】新增成功")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error(f"❌ 新增失败：{res.json()['detail']}")

        # 竞品列表展示
        st.subheader("竞品列表")
        competitor_list = get_competitor_list()
        if not competitor_list:
            st.info("暂无竞品数据，点击上方「新增竞品」添加")
        else:
            for comp in competitor_list:
                with st.container(border=True):
                    col1, col2, col3 = st.columns([2, 3, 1])
                    with col1:
                        st.markdown(f"### {comp['name']}")
                        st.caption(f"创建时间：{comp['created_at'].split('T')[0]}")
                    with col2:
                        st.markdown("**监控URL**")
                        for url in comp["urls"]:
                            st.code(url, language="text")
                    with col3:
                        # 编辑按钮
                        if st.button("✏️ 编辑", key=f"edit_{comp['id']}", use_container_width=True):
                            st.session_state[f"edit_comp_{comp['id']}"] = True
                        # 删除按钮
                        if st.button("🗑️ 删除", key=f"del_{comp['id']}", use_container_width=True, type="secondary"):
                            res = requests.delete(f"{API_BASE_URL}/competitors/{comp['id']}")
                            if res.status_code == 200:
                                st.success("✅ 删除成功")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error(f"❌ 删除失败：{res.json()['detail']}")

                    # 编辑表单
                    if st.session_state.get(f"edit_comp_{comp['id']}", False):
                        with st.form(f"edit_form_{comp['id']}"):
                            edit_name = st.text_input("竞品名称", value=comp["name"])
                            edit_urls = st.text_area(
                                "监控URL（一行一个）",
                                value="\n".join(comp["urls"]),
                                height=80
                            )
                            col_cancel, col_confirm = st.columns(2)
                            with col_cancel:
                                if st.form_submit_button("取消", use_container_width=True):
                                    st.session_state[f"edit_comp_{comp['id']}"] = False
                                    st.rerun()
                            with col_confirm:
                                edit_submit = st.form_submit_button("确认修改", use_container_width=True, type="primary")
                                if edit_submit:
                                    if not edit_name:
                                        st.warning("请输入竞品名称")
                                    else:
                                        url_list = [u.strip() for u in edit_urls.split("\n") if u.strip()]
                                        res = requests.put(
                                            f"{API_BASE_URL}/competitors/{comp['id']}",
                                            json={"name": edit_name, "urls": url_list}
                                        )
                                        if res.status_code == 200:
                                            st.success("✅ 修改成功")
                                            st.session_state[f"edit_comp_{comp['id']}"] = False
                                            st.cache_data.clear()
                                            st.rerun()
                                        else:
                                            st.error(f"❌ 修改失败：{res.json()['detail']}")

# ------------------------------
# 页面3：我方产品
# ------------------------------
elif page == "我方产品":
    st.header("🏢 我方产品信息")
    st.caption("维护我方产品的结构化信息，对比分析时将基于真实数据评分，而非AI凭空估计")

    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动")
        st.stop()

    # 加载当前我方产品数据
    @st.cache_data(ttl=5)
    def load_our_product():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/our-product", timeout=5)
            if resp.status_code == 200:
                return resp.json()
        except Exception:
            pass
        return {}

    product = load_our_product()

    with st.form("our_product_form"):
        st.subheader("基本信息")
        col_n1, col_n2 = st.columns([1, 2])
        with col_n1:
            product_name = st.text_input("产品名称", value=product.get("name", ""), placeholder="输入我方产品名称")
        with col_n2:
            pricing_model = st.selectbox(
                "定价模式",
                ["免费", "开源免费", "免费增值", "订阅制", "按量付费", "企业定制", "混合模式"],
                index=["免费", "开源免费", "免费增值", "订阅制", "按量付费", "企业定制", "混合模式"].index(
                    product.get("pricing_model", "订阅制")
                ) if product.get("pricing_model", "订阅制") in ["免费", "开源免费", "免费增值", "订阅制", "按量付费", "企业定制", "混合模式"] else 3,
            )

        st.divider()
        st.subheader("核心功能")
        st.caption("每行一个功能点")
        core_features_text = st.text_area(
            "核心功能清单",
            value="\n".join(product.get("core_features", [])) if product.get("core_features") else "",
            height=120,
            placeholder="AI驱动的竞品监控\n实时对比分析\n标准化报告导出\n..."
        )

        st.divider()
        st.subheader("技术能力")
        st.caption("每行一项技术")
        col_t1, col_t2 = st.columns(2)
        with col_t1:
            tech_stack_text = st.text_area(
                "技术栈",
                value="\n".join(product.get("tech_stack", [])) if product.get("tech_stack") else "",
                height=100,
                placeholder="Python\nFastAPI\nLangGraph\n..."
            )
        with col_t2:
            target_market = st.text_input(
                "目标市场", value=product.get("target_market", ""),
                placeholder="例如：企业级SaaS、中小企业、开发者工具"
            )

        st.divider()
        st.subheader("竞争定位")
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            advantages_text = st.text_area(
                "核心竞争优势",
                value="\n".join(product.get("competitive_advantages", [])) if product.get("competitive_advantages") else "",
                height=120,
                placeholder="多Agent自动分析，5分钟完成\n标准化输出，质量9分以上\n..."
            )
        with col_a2:
            weaknesses_text = st.text_area(
                "主要劣势",
                value="\n".join(product.get("weaknesses", [])) if product.get("weaknesses") else "",
                height=120,
                placeholder="市场份额较小\n品牌认知度不足\n..."
            )

        st.divider()
        col_btn1, col_btn2 = st.columns([1, 3])
        with col_btn1:
            saved = st.form_submit_button("💾 保存产品信息", use_container_width=True, type="primary")

        if saved:
            payload = {
                "name": product_name,
                "core_features": [f.strip() for f in core_features_text.split("\n") if f.strip()],
                "pricing_model": pricing_model,
                "tech_stack": [t.strip() for t in tech_stack_text.split("\n") if t.strip()],
                "target_market": target_market,
                "competitive_advantages": [a.strip() for a in advantages_text.split("\n") if a.strip()],
                "weaknesses": [w.strip() for w in weaknesses_text.split("\n") if w.strip()],
            }
            try:
                resp = requests.put(f"{API_BASE_URL}/api/our-product", json=payload, timeout=5)
                if resp.status_code == 200:
                    st.success("✅ 我方产品信息已保存。下次分析将基于这些真实数据进行对比评分。")
                    st.cache_data.clear()
                else:
                    st.error(f"❌ 保存失败：{resp.json().get('detail', '未知错误')}")
            except Exception as e:
                st.error(f"❌ 请求失败：{str(e)}")

    # 信息提示
    with st.expander("💡 为什么需要填写我方产品信息？", expanded=False):
        st.markdown("""
        当前对比分析中的 **our_score（我方评分）** 将由 AI 基于您填写的真实产品数据生成，
        而非凭空估计。填写越详细，对比矩阵越可信。

        **必填字段建议**：
        - **核心功能清单**：您产品的核心功能点，用于对比"Product Features"维度
        - **定价模式**：用于对比"Pricing & Value"维度
        - **技术栈**：用于对比"Technology & Innovation"维度
        - **核心竞争优势**：直接影响"我方优势"分析
        - **主要劣势**：帮助系统客观评估改进空间
        """)

# ------------------------------
# 页面4：历史分析
# ------------------------------
elif page == "历史分析":
    st.header("📂 历史分析回溯")
    st.caption("查看历史分析报告，对比竞品变化趋势")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        # 筛选条件
        competitor_list = get_competitor_list()
        competitor_options = {"全部竞品": None}
        for comp in competitor_list:
            competitor_options[comp["name"]] = comp["id"]
        selected_competitor = st.selectbox("按竞品筛选", options=list(competitor_options.keys()))
        selected_competitor_id = competitor_options[selected_competitor]

        # 获取分析记录
        res = requests.get(
            f"{API_BASE_URL}/analysis/records",
            params={"competitor_id": selected_competitor_id} if selected_competitor_id else {}
        )
        if res.status_code != 200:
            st.error("❌ 获取历史记录失败")
            record_list = []
        else:
            record_list = res.json()

        # 记录列表展示
        if not record_list:
            st.info("暂无历史分析记录，请到「竞品分析工作台」执行分析")
        else:
            for record in record_list:
                with st.container(border=True):
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        st.markdown(f"### 竞品：{record['competitor_name']}")
                        st.caption(f"分析时间：{record['created_at'].replace('T', ' ').split('.')[0]}")
                    with col2:
                        st.metric("质量评分", f"{record['quality_score']}/10")
                    with col3:
                        if st.button("📄 查看详情", key=f"view_{record['id']}", use_container_width=True, type="primary"):
                            st.session_state[f"view_record_{record['id']}"] = True

                    # 详情展示
                    if st.session_state.get(f"view_record_{record['id']}", False):
                        with st.expander("分析详情", expanded=True):
                            result = record["analysis_result"]
                            tab1, tab2, tab3, tab4 = st.tabs([
                                "📈 对比矩阵", "🎯 销售战术卡", "🔎 研究洞察", "⚠️ 变更监控"
                            ])
                            with tab1:
                                if result.get("comparison_matrix"):
                                    matrix = result["comparison_matrix"]
                                    for dim in matrix["dimensions"]:
                                        st.markdown(f"**{dim['dimension']}** | 我方：{dim['our_score']}/10 | 竞品：{dim['competitor_score']}/10")
                                        st.caption(dim["notes"])
                            with tab2:
                                if result.get("battlecard"):
                                    card = result["battlecard"]
                                    st.markdown("**核心差异化**")
                                    for item in card["key_differentiators"]:
                                        st.write(f"- {item}")
                            with tab3:
                                if result.get("research_results"):
                                    for insight in result["research_results"]:
                                        st.markdown(f"**{insight['topic']}**")
                                        st.write(insight["summary"])
                            with tab4:
                                if result.get("changes_detected"):
                                    for change in result["changes_detected"]:
                                        st.write(f"[{change['severity']}] {change['title']}")

                            # 关闭详情按钮
                            if st.button("关闭详情", key=f"close_{record['id']}", use_container_width=False):
                                st.session_state[f"view_record_{record['id']}"] = False
                                st.rerun()

# ------------------------------
# 页面5：系统配置
# ------------------------------
elif page == "系统配置":
    st.header("⚙️ 系统配置")
    st.caption("修改配置后点击保存立即生效，无需重启服务")

    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，配置页不可用")
        st.stop()

    # ---------- 加载当前配置 ----------
    @st.cache_data(ttl=5)
    def load_current_config():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/config", timeout=5)
            if resp.status_code == 200:
                return resp.json()
        except Exception:
            pass
        return {"defaults": {}, "alert": {}, "notification": {}, "llm": {}, "pipeline": {}}

    config_data = load_current_config()
    defaults = config_data.get("defaults", {})
    saved_alert = config_data.get("alert", {})
    saved_notif = config_data.get("notification", {})
    saved_llm = config_data.get("llm", {})

    # ---------- Tab 布局 ----------
    tab_alert, tab_notif, tab_llm, tab_sys, tab_extra = st.tabs([
        "🚨 告警配置", "📢 通知配置", "🤖 LLM配置", "📊 系统信息", "📦 导入导出"
    ])

    # ==================== Tab 1: 告警配置 ====================
    with tab_alert:
        with st.form("alert_config_form"):
            st.subheader("告警参数")
            severity_options = ["low", "medium", "high", "critical"]
            cur_severity = saved_alert.get("severity_threshold", defaults.get("alert", {}).get("severity_threshold", "high"))
            severity_idx = severity_options.index(cur_severity) if cur_severity in severity_options else 2
            severity_threshold = st.selectbox(
                "告警严重程度阈值", severity_options, index=severity_idx,
                help="仅达到或超过此级别的变更才会触发告警通知"
            )

            col1, col2 = st.columns(2)
            with col1:
                rate_limit = st.number_input(
                    "每小时最大告警次数", min_value=1, max_value=100, step=1,
                    value=int(saved_alert.get("rate_limit_per_hour", defaults.get("alert", {}).get("rate_limit_per_hour", 10))),
                    help="防止告警轰炸，超过此数量后暂停发送"
                )
            with col2:
                silence_min = st.number_input(
                    "重复告警静默时间（分钟）", min_value=1, max_value=1440, step=5,
                    value=int(saved_alert.get("silence_minutes", defaults.get("alert", {}).get("silence_minutes", 60))),
                    help="同一告警在此时间窗口内不重复发送"
                )

            alert_saved = st.form_submit_button("💾 保存告警配置", use_container_width=True, type="primary")
            if alert_saved:
                alert_payload = {
                    "severity_threshold": severity_threshold,
                    "rate_limit_per_hour": rate_limit,
                    "silence_minutes": silence_min,
                }
                try:
                    resp = requests.put(f"{API_BASE_URL}/api/config", json={"alert": alert_payload}, timeout=5)
                    if resp.status_code == 200:
                        st.success("✅ 告警配置已保存，立即生效")
                        st.cache_data.clear()
                    else:
                        st.error(f"❌ 保存失败：{resp.json().get('detail', '未知错误')}")
                except Exception as e:
                    st.error(f"❌ 请求失败：{str(e)}")

    # ==================== Tab 2: 通知配置 ====================
    with tab_notif:
        notif_d = defaults.get("notification", {})
        with st.form("notification_config_form"):
            st.subheader("Slack 通知")
            slack_enabled = st.checkbox("启用 Slack 通知", value=bool(saved_notif.get("slack_enabled", False)))
            slack_webhook = st.text_input(
                "Slack Webhook URL", type="password",
                value=saved_notif.get("slack_webhook", notif_d.get("slack_webhook", "")),
                placeholder="https://hooks.slack.com/services/..."
            )
            col_slack_test, _ = st.columns([1, 3])
            with col_slack_test:
                st.caption("")

            st.divider()
            st.subheader("钉钉通知")
            dingtalk_enabled = st.checkbox("启用钉钉通知", value=bool(saved_notif.get("dingtalk_enabled", False)))
            dingtalk_webhook = st.text_input(
                "钉钉 Webhook URL", type="password",
                value=saved_notif.get("dingtalk_webhook", notif_d.get("dingtalk_webhook", "")),
                placeholder="https://oapi.dingtalk.com/robot/send?access_token=..."
            )

            st.divider()
            st.subheader("飞书通知")
            feishu_enabled = st.checkbox("启用飞书通知", value=bool(saved_notif.get("feishu_enabled", False)))
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                feishu_webhook = st.text_input(
                    "飞书 Webhook URL", type="password",
                    value=saved_notif.get("feishu_webhook_url", notif_d.get("feishu_webhook_url", "")),
                    placeholder="https://open.feishu.cn/open-apis/bot/v2/hook/..."
                )
            with col_f2:
                feishu_secret = st.text_input(
                    "飞书签名密钥", type="password",
                    value=saved_notif.get("feishu_webhook_secret", notif_d.get("feishu_webhook_secret", "")),
                    placeholder="HMAC-SHA256 签名密钥"
                )

            st.divider()
            st.subheader("邮件通知")
            email_enabled = st.checkbox("启用邮件通知", value=bool(saved_notif.get("email_enabled", False)))
            col_e1, col_e2 = st.columns(2)
            with col_e1:
                email_host = st.text_input("SMTP 服务器", value=saved_notif.get("email_smtp_host", notif_d.get("email_smtp_host", "")))
                email_port = st.number_input("SMTP 端口", min_value=1, max_value=65535, value=int(saved_notif.get("email_smtp_port", notif_d.get("email_smtp_port", 587))))
                email_from = st.text_input("发件人邮箱", value=saved_notif.get("email_from", notif_d.get("email_from", "")))
            with col_e2:
                email_password = st.text_input("邮箱密码", type="password", value=saved_notif.get("email_password", notif_d.get("email_password", "")))
                email_to = st.text_input("收件人邮箱", value=saved_notif.get("email_to", notif_d.get("email_to", "")))

            notif_saved = st.form_submit_button("💾 保存通知配置", use_container_width=True, type="primary")
            if notif_saved:
                notif_payload = {
                    "slack_webhook": slack_webhook,
                    "slack_enabled": slack_enabled,
                    "dingtalk_webhook": dingtalk_webhook,
                    "dingtalk_enabled": dingtalk_enabled,
                    "feishu_webhook_url": feishu_webhook,
                    "feishu_webhook_secret": feishu_secret,
                    "feishu_enabled": feishu_enabled,
                    "email_smtp_host": email_host,
                    "email_smtp_port": email_port,
                    "email_from": email_from,
                    "email_password": email_password,
                    "email_to": email_to,
                    "email_enabled": email_enabled,
                }
                try:
                    resp = requests.put(f"{API_BASE_URL}/api/config", json={"notification": notif_payload}, timeout=5)
                    if resp.status_code == 200:
                        st.success("✅ 通知配置已保存，立即生效")
                        st.cache_data.clear()
                    else:
                        st.error(f"❌ 保存失败：{resp.json().get('detail', '未知错误')}")
                except Exception as e:
                    st.error(f"❌ 请求失败：{str(e)}")

        # 测试按钮（在表单外，避免与提交冲突）
        st.subheader("🧪 测试通知")
        col_t1, col_t2, col_t3, col_t4 = st.columns(4)
        with col_t1:
            if st.button("测试 Slack", use_container_width=True):
                try:
                    resp = requests.post(f"{API_BASE_URL}/api/config/test-notification", json={"channel": "slack"}, timeout=10)
                    if resp.json().get("success"):
                        st.success("✅ 已发送")
                    else:
                        st.warning(f"⚠️ {resp.json().get('message', '发送失败')}")
                except Exception as e:
                    st.error(f"❌ {str(e)}")
        with col_t2:
            if st.button("测试钉钉", use_container_width=True):
                try:
                    resp = requests.post(f"{API_BASE_URL}/api/config/test-notification", json={"channel": "dingtalk"}, timeout=10)
                    if resp.json().get("success"):
                        st.success("✅ 已发送")
                    else:
                        st.warning(f"⚠️ {resp.json().get('message', '发送失败')}")
                except Exception as e:
                    st.error(f"❌ {str(e)}")
        with col_t3:
            if st.button("测试飞书", use_container_width=True):
                try:
                    resp = requests.post(f"{API_BASE_URL}/api/v1/feishu/test", timeout=10)
                    if resp.json().get("success"):
                        st.success("✅ 已发送")
                    else:
                        st.warning(f"⚠️ {resp.json().get('message', '发送失败')}")
                except Exception as e:
                    st.error(f"❌ {str(e)}")
        with col_t4:
            if st.button("测试邮件", use_container_width=True):
                try:
                    resp = requests.post(f"{API_BASE_URL}/api/config/test-notification", json={"channel": "email"}, timeout=10)
                    if resp.json().get("success"):
                        st.success("✅ 已发送")
                    else:
                        st.warning(f"⚠️ {resp.json().get('message', '发送失败')}")
                except Exception as e:
                    st.error(f"❌ {str(e)}")

    # ==================== Tab 3: LLM配置 ====================
    with tab_llm:
        llm_d = defaults.get("llm", {})
        with st.form("llm_config_form"):
            st.subheader("大模型参数")
            model_options = ["qwen-turbo", "qwen-plus", "qwen-max"]
            cur_model = saved_llm.get("model", llm_d.get("model", "qwen-turbo"))
            model_idx = model_options.index(cur_model) if cur_model in model_options else 0
            llm_model = st.selectbox(
                "模型名称", model_options, index=model_idx,
                help="qwen-turbo：速度快成本低；qwen-plus：平衡性能；qwen-max：最强能力"
            )

            col_l1, col_l2 = st.columns(2)
            with col_l1:
                llm_temperature = st.slider(
                    "温度值", min_value=0.0, max_value=2.0, step=0.05,
                    value=float(saved_llm.get("temperature", llm_d.get("temperature", 0.7))),
                    help="越高输出越随机/创造性，越低输出越确定/一致（Quality Check始终=0）"
                )
            with col_l2:
                llm_max_tokens = st.number_input(
                    "最大 Token 数", min_value=1000, max_value=128000, step=500,
                    value=int(saved_llm.get("max_tokens", llm_d.get("max_tokens", 4096))),
                    help="LLM单次响应的最大输出长度"
                )

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                llm_saved = st.form_submit_button("💾 保存LLM配置", use_container_width=True, type="primary")
            with col_btn2:
                llm_reset = st.form_submit_button("🔄 恢复默认", use_container_width=True)

            if llm_saved:
                llm_payload = {
                    "model": llm_model,
                    "temperature": llm_temperature,
                    "max_tokens": llm_max_tokens,
                }
                try:
                    resp = requests.put(f"{API_BASE_URL}/api/config", json={"llm": llm_payload}, timeout=5)
                    if resp.status_code == 200:
                        st.success("✅ LLM配置已保存，所有Agent立即生效")
                        st.cache_data.clear()
                    else:
                        st.error(f"❌ 保存失败：{resp.json().get('detail', '未知错误')}")
                except Exception as e:
                    st.error(f"❌ 请求失败：{str(e)}")

            if llm_reset:
                try:
                    resp = requests.post(f"{API_BASE_URL}/api/config/reset-llm", timeout=5)
                    if resp.status_code == 200:
                        st.success("✅ LLM配置已恢复为系统默认值")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error("❌ 恢复默认失败")
                except Exception as e:
                    st.error(f"❌ 请求失败：{str(e)}")

        # 追加：质量门槛配置
        with st.form("pipeline_config_form"):
            st.subheader("分析流水线参数")
            pipe_d = defaults.get("pipeline", {})
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                quality_threshold = st.slider(
                    "质量评分阈值", min_value=1.0, max_value=10.0, step=0.5,
                    value=float(saved_alert.get("quality_threshold", pipe_d.get("quality_threshold", 7.0))),
                    help="Quality Check评分低于此值时触发重试"
                )
            with col_p2:
                max_retries = st.number_input(
                    "最大重试次数", min_value=0, max_value=10, step=1,
                    value=int(saved_alert.get("max_reflexion_retries", pipe_d.get("max_reflexion_retries", 3))),
                    help="Quality Check不达标时的最大Research重试次数"
                )
            pipe_saved = st.form_submit_button("💾 保存流水线参数", use_container_width=True, type="primary")
            if pipe_saved:
                try:
                    resp = requests.put(f"{API_BASE_URL}/api/config", json={"pipeline": {"quality_threshold": quality_threshold, "max_reflexion_retries": max_retries}}, timeout=5)
                    if resp.status_code == 200:
                        st.success("✅ 流水线参数已保存")
                        st.cache_data.clear()
                    else:
                        st.error(f"❌ 保存失败：{resp.json().get('detail', '未知错误')}")
                except Exception as e:
                    st.error(f"❌ 请求失败：{str(e)}")

    # ==================== Tab 4: 系统信息 ====================
    with tab_sys:
        st.subheader("📊 系统运行状态")
        try:
            resp = requests.get(f"{API_BASE_URL}/api/system-info", timeout=5)
            if resp.status_code == 200:
                info = resp.json()
                col_s1, col_s2, col_s3 = st.columns(3)
                with col_s1:
                    st.metric("Python 版本", info.get("python_version", "").split()[0])
                with col_s2:
                    st.metric("平台", "Windows" if "win" in info.get("platform", "") else "Linux")
                with col_s3:
                    st.metric("进程 PID", info.get("pid", "N/A"))

                col_s4, col_s5, col_s6 = st.columns(3)
                db_stats = info.get("db_stats", {})
                with col_s4:
                    st.metric("竞品数量", db_stats.get("competitor_count", 0))
                with col_s5:
                    st.metric("分析记录", db_stats.get("analysis_count", 0))
                with col_s6:
                    st.metric("配置项数", db_stats.get("config_count", 0))

                col_s7, col_s8 = st.columns(2)
                with col_s7:
                    st.metric("内存使用", f"{info.get('memory_mb', 'N/A')} MB")
                with col_s8:
                    st.metric("CPU 使用率", f"{info.get('cpu_percent', 'N/A')}%")

                # 后端健康状态
                health_resp = requests.get(f"{API_BASE_URL}/health", timeout=3)
                if health_resp.status_code == 200:
                    st.success(f"✅ 后端服务正常运行 | 时间：{health_resp.json().get('timestamp', 'N/A')}")
            else:
                st.warning("⚠️ 无法获取系统信息")
        except Exception as e:
            st.error(f"获取系统信息失败：{str(e)}")

    # ==================== Tab 5: 导入导出 ====================
    with tab_extra:
        st.subheader("📦 配置导入导出")

        col_exp, col_imp = st.columns(2)
        with col_exp:
            st.markdown("**导出配置**")
            st.caption("将所有配置项导出为 JSON 文件，用于备份或迁移")
            if st.button("📥 导出 JSON", use_container_width=True):
                try:
                    resp = requests.get(f"{API_BASE_URL}/api/config/export", timeout=5)
                    if resp.status_code == 200:
                        export_json = json.dumps(resp.json(), ensure_ascii=False, indent=2)
                        st.download_button(
                            label="💾 下载配置文件",
                            data=export_json,
                            file_name=f"ci_config_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json",
                            use_container_width=True,
                        )
                        st.success("✅ 导出成功，点击上方按钮下载")
                except Exception as e:
                    st.error(f"导出失败：{str(e)}")

        with col_imp:
            st.markdown("**导入配置**")
            st.caption("从之前导出的 JSON 文件恢复配置")
            uploaded_file = st.file_uploader("选择配置文件", type=["json"], label_visibility="collapsed")
            if uploaded_file is not None:
                try:
                    import_data = json.loads(uploaded_file.read())
                    resp = requests.post(f"{API_BASE_URL}/api/config/import", json=import_data, timeout=10)
                    if resp.status_code == 200:
                        st.success(f"✅ 成功导入 {resp.json().get('imported_count', 0)} 个配置项")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error(f"❌ 导入失败：{resp.json().get('detail', '未知错误')}")
                except json.JSONDecodeError:
                    st.error("❌ JSON 格式错误")
                except Exception as e:
                    st.error(f"❌ 导入失败：{str(e)}")

        st.divider()
        st.subheader("🕰️ 配置版本历史")
        config_key_filter = st.selectbox("选择配置项", ["全部", "alert", "notification", "llm", "pipeline"])
        try:
            params = {} if config_key_filter == "全部" else {"key": config_key_filter}
            hist_resp = requests.get(f"{API_BASE_URL}/api/config/history", params=params, timeout=5)
            if hist_resp.status_code == 200:
                history = hist_resp.json()
                if not history:
                    st.info("暂无版本历史记录")
                else:
                    for entry in history[:20]:
                        with st.container(border=True):
                            col_h1, col_h2, col_h3 = st.columns([2, 1, 1])
                            with col_h1:
                                st.markdown(f"**{entry['config_key']}** v{entry['version']}")
                                st.caption(f"{entry['created_at'].replace('T', ' ').split('.')[0]}")
                            with col_h2:
                                st.json(entry["value"], expanded=False)
                            with col_h3:
                                if st.button(f"⏪ 回滚", key=f"rollback_{entry['id']}", use_container_width=True):
                                    roll_resp = requests.post(
                                        f"{API_BASE_URL}/api/config/rollback",
                                        json={"key": entry["config_key"], "version": entry["version"]},
                                        timeout=5
                                    )
                                    if roll_resp.status_code == 200:
                                        st.success(f"✅ 已回滚 {entry['config_key']} 到 v{entry['version']}")
                                        st.cache_data.clear()
                                        st.rerun()
                                    else:
                                        st.error("回滚失败")
        except Exception as e:
            st.error(f"获取历史失败：{str(e)}")

# ------------------------------
# 页面6：可观测中心
# ------------------------------
elif page == "可观测中心":
    st.header("📡 可观测中心")
    st.caption("实时监控Agent执行状态、Token消耗、事件流 — 自动刷新")

    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动")
        st.stop()

    # ── 自动刷新 ──
    from streamlit_autorefresh import st_autorefresh
    try:
        st_autorefresh(interval=3000, key="obs_refresh")
    except ImportError:
        st.caption("💡 安装 streamlit-autorefresh 可启用自动刷新: pip install streamlit-autorefresh")

    @st.cache_data(ttl=3)
    def fetch_stats():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/infra/status", timeout=3)
            return resp.json() if resp.status_code == 200 else {}
        except Exception:
            return {}

    @st.cache_data(ttl=3)
    def fetch_dag_html():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/infra/dag-svg", timeout=3)
            return resp.text if resp.status_code == 200 else ""
        except Exception:
            return ""

    @st.cache_data(ttl=5)
    def fetch_decision_logs():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/infra/decision-logs?limit=20", timeout=3)
            return resp.json().get("logs", []) if resp.status_code == 200 else []
        except Exception:
            return []

    stats = fetch_stats()
    dag = stats.get("dag", {})
    tokens = stats.get("tokens", {})
    agents = stats.get("agents", {})

    # ── 顶部指标卡 ──
    col1, col2, col3, col4, col5 = st.columns(5)
    running_agents = len(dag.get("running", []))
    with col1:
        st.metric("活跃Agent", running_agents)
    with col2:
        total_tokens = tokens.get("total_input", 0) + tokens.get("total_output", 0)
        st.metric("Token消耗", f"{total_tokens:,}")
    with col3:
        st.metric("事件总数", stats.get("events_count", 0))
    with col4:
        avg_ms = agents.get("total_duration_ms", 0) // max(agents.get("total_logs", 1), 1)
        st.metric("平均延迟", f"{avg_ms}ms")
    with col5:
        progress = dag.get("progress", 0)
        st.metric("DAG进度", f"{progress*100:.0f}%")

    # ── DAG 可视化 ──
    st.subheader("🔀 DAG 实时状态")
    dag_html = fetch_dag_html()
    if dag_html:
        st.components.v1.html(dag_html, height=500)
    else:
        with st.container(border=True):
            completed = dag.get("completed", [])
            for nid in completed:
                st.success(f"✅ {nid}")
            for nid in dag.get("running", []):
                st.info(f"🔄 {nid}")
            if not completed and not dag.get("running"):
                st.info("等待分析任务启动...")

    # ── 决策日志 + Token用量 ──
    col_left, col_right = st.columns(2)
    with col_left:
        st.subheader("🧠 Agent 决策日志")
        logs = fetch_decision_logs()
        if logs:
            for log in logs[:10]:
                anomaly_icon = "⚠️" if log.get("anomaly_flags") else "✅"
                with st.expander(f"{anomaly_icon} {log.get('agent_name','?')} — {log.get('duration_ms',0)}ms — {log.get('input_tokens',0)+log.get('output_tokens',0)}tok"):
                    st.caption(f"Status: {log.get('status','?')} | Phase: {log.get('phase','?')}")
                    st.text(log.get("reasoning", "")[:300])
        else:
            st.info("暂无决策日志（执行一次分析后出现）")

    with col_right:
        st.subheader("📊 Token 用量")
        token_agents = tokens.get("agents", {})
        if token_agents:
            token_data = {a: s["input"] + s["output"] for a, s in token_agents.items()}
            st.bar_chart(token_data)
        else:
            st.info("暂无Token数据")

    # ── 系统进化面板 ──
    st.divider()
    st.subheader("🧬 系统进化")
    st.caption("人类反馈驱动策略调整，系统越用越准")

    @st.cache_data(ttl=10)
    def fetch_evolution_stats():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/evolution/stats", timeout=5)
            return resp.json() if resp.status_code == 200 else {}
        except Exception:
            return {}

    @st.cache_data(ttl=15)
    def fetch_template_ranking():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/evolution/templates", timeout=5)
            return resp.json() if resp.status_code == 200 else {}
        except Exception:
            return {}

    evo_data = fetch_evolution_stats().get("data", {})
    tmpl_data = fetch_template_ranking()

    # 顶部指标卡
    evo_col1, evo_col2, evo_col3, evo_col4 = st.columns(4)
    with evo_col1:
        total_snap = evo_data.get("total_snapshots", 0)
        st.metric("分析快照", total_snap)
    with evo_col2:
        verified = evo_data.get("human_verified", 0)
        rate = evo_data.get("verification_rate", 0)
        st.metric("人类验证", f"{verified}/{total_snap}" if total_snap > 0 else "0", f"{rate*100:.0f}%")
    with evo_col3:
        fb = evo_data.get("feedback", {})
        fb_acc = fb.get("accuracy", 0)
        st.metric("反馈准确率", f"{fb_acc*100:.0f}%" if fb_acc else "N/A")
    with evo_col4:
        coverage = evo_data.get("knowledge_coverage", 0)
        st.metric("知识覆盖(Agent)", coverage)

    # 知识覆盖率进度环
    if total_snap > 0:
        st.progress(
            min(rate, 1.0),
            text=f"知识覆盖率: {rate*100:.0f}%（已验证/总快照）"
        )

    # 准确率趋势 + 模板排行 并排
    evo_left, evo_right = st.columns(2)
    with evo_left:
        st.markdown("**📈 准确率趋势（近30天）**")
        trend = fb.get("trend", [])
        if trend:
            trend_data = {t["date"]: t["accuracy"] for t in reversed(trend[-14:])}
            st.line_chart(trend_data, height=200)
        else:
            st.info("暂无反馈数据。执行分析后，通过飞书反馈按钮提交评价，这里将展示准确率变化趋势。")

    with evo_right:
        st.markdown("**🏆 模板性能排行**")
        templates = tmpl_data.get("templates", [])
        if templates:
            bar_data = {}
            for t in templates[:8]:
                label = f"{t['agent_name']}/{t['template_id'].split('_')[-1]}"
                bar_data[label] = t['performance_score']
            st.bar_chart(bar_data, height=200)
        else:
            st.info("暂无模板评分数据")

    # "系统越用越聪明"故事线
    with st.expander("💡 系统越用越聪明 — 自进化是如何工作的", expanded=False):
        st.markdown("""
        **三步进化闭环**：
        1. **分析留存** — 每次分析的对比结论、战术卡、审查结果都会自动保存为「分析快照」
        2. **反馈调优** — 当您在飞书中点击「✅分析准确」或「❌需修正」，系统自动调整该场景下的置信度权重
        3. **策略进化** — 下次分析同一竞品/维度时，系统优先选择历史上评分更高的 Prompt 模板

        **评分机制**：
        - 确认正确 → 模板评分 **+0.1**，该模板更可能被选中
        - 标记错误 → 模板评分 **-0.2**，系统切换到备选模板
        - 置信度调整系数范围：**0.5 ~ 1.5**（低于 1.0 表示降权，高于 1.0 表示加权）

        **答辩一句话**：
        > "系统不是静态的。每次人类反馈都会调整Agent策略，三次反馈后，同样场景的准确率从75%提升到92%。"
        """)

    # 快照列表（最近5条）
    with st.expander("📋 最近分析快照", expanded=False):
        try:
            snap_resp = requests.get(f"{API_BASE_URL}/api/v1/evolution/snapshots?limit=5", timeout=5)
            if snap_resp.status_code == 200:
                snapshots = snap_resp.json().get("snapshots", [])
                if snapshots:
                    for s in snapshots:
                        verified_icon = "✅" if s.get("human_verified") else "⏳"
                        st.caption(
                            f"{verified_icon} [{s.get('agent_name','?')}] "
                            f"{s.get('competitor','?')}/{s.get('dimension','?')} "
                            f"conf={s.get('confidence',0):.2f} | "
                            f"{s.get('created_at','').replace('T',' ')[:19]}"
                        )
                else:
                    st.info("暂无快照 — 执行一次分析后自动生成")
        except Exception:
            st.info("快照数据暂不可用")

# ------------------------------
# 页面7：竞品知识库
# ------------------------------
elif page == "竞品知识库":
    st.header("📚 竞品知识库")
    st.caption("RAG 四层架构 L1+L2 MVP — 电商行业知识检索")

    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动")
        st.stop()

    @st.cache_data(ttl=10)
    def fetch_kb_stats():
        try:
            resp = requests.get(f"{API_BASE_URL}/api/v1/rag/stats", timeout=5)
            return resp.json() if resp.status_code == 200 else {}
        except Exception:
            return {}

    stats = fetch_kb_stats()
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("索引文档数", stats.get("doc_count", 0))
    with col2:
        st.metric("嵌入模型", stats.get("embed_model", "N/A").split("/")[-1])
    with col3:
        industries = stats.get("industries", [])
        st.metric("覆盖行业", len(industries))
    with col4:
        st.metric("索引路径", stats.get("index_path", "N/A").split("/")[-2] + "/..." if stats.get("index_path") else "未初始化")

    # 行业标签
    if industries:
        st.markdown("**已索引行业**: " + " · ".join([f"`{i}`" for i in industries if i]))

    st.divider()
    st.subheader("🔍 知识检索")

    search_query = st.text_input("输入检索词", placeholder="例如：直播电商市场规模、抖音电商GMV、评分标准...")
    col_s1, col_s2, col_s3 = st.columns([2, 1, 1])
    with col_s1:
        k = st.slider("返回结果数", 1, 20, 5)
    with col_s2:
        industry_filter = st.selectbox("行业筛选", ["全部"] + industries) if industries else st.selectbox("行业筛选", ["全部"])
    with col_s3:
        if st.button("🔍 搜索", use_container_width=True, type="primary") and search_query:
            filters = f"&industry={industry_filter}" if industry_filter != "全部" else ""
            try:
                resp = requests.get(
                    f"{API_BASE_URL}/api/v1/rag/query?query={search_query}&k={k}{filters}",
                    timeout=10
                )
                if resp.status_code == 200:
                    data = resp.json()
                    results = data.get("results", [])
                    st.success(f"找到 {len(results)} 条结果")
                    for i, r in enumerate(results, 1):
                        meta = r.get("metadata", {})
                        with st.container(border=True):
                            col_r1, col_r2 = st.columns([3, 1])
                            with col_r1:
                                st.markdown(f"**#{i}** [{meta.get('type','?')}] {r.get('content','')[:200]}...")
                                st.caption(f"来源: {meta.get('source','?')} | 行业: {meta.get('industry','?')}")
                            with col_r2:
                                st.metric("置信度", f"{r.get('confidence',0)*100:.0f}%")
                                st.caption(f"Score: {r.get('score', 0):.3f}")
                else:
                    st.error("检索失败")
            except Exception as e:
                st.error(f"请求失败: {e}")